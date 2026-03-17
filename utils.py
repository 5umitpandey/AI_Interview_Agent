#utils.py

import pdfplumber
import docx
import json
from openai import AsyncOpenAI
import os
from dotenv import load_dotenv

load_dotenv()

# Initialize OpenAI client using environment variable
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise RuntimeError("OPENAI_API_KEY not set in environment")

client = AsyncOpenAI(api_key=OPENAI_API_KEY)


async def prepare_interview_plan(resume_text: str, job_description: str | None = None) -> dict:
    """
    Generate an interview plan with questions based on the candidate's resume.
    
    Args:
        resume_text: Extracted text from the candidate's resume
        
    Returns:
        dict: Interview plan with summary and questions
    """
    
        # Short JD-aware prompt: produce a 2-sentence summary and 10 concise questions
    jd_block = f"Job Description:\n{job_description}\n\n" if job_description else ""

    prompt = (
        jd_block
        + "Given the candidate resume below, return a JSON object with: 'summary' (2 sentences) and 'questions' (exactly 3 concise, JD-relevant interview questions).\n"
        + "Rules: questions <= 30 words, mix technical/behavioral/problem-solving, prioritize JD skills when provided. Respond with JSON only.\n\n"
        + f"Resume:\n{resume_text}\n"
    )
    
    try:
        response = await client.chat.completions.create(
            model="gpt-4.1-nano",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert technical interviewer and HR professional. Produce concise, targeted interview questions and a short summary. Reply with valid JSON only."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.25,
            max_tokens=1000,
        )

        text = response.choices[0].message.content.strip()
        usage = response.usage

        prompt_tokens = usage.prompt_tokens if usage else 0
        completion_tokens = usage.completion_tokens if usage else 0
        total_tokens = usage.total_tokens if usage else 0

        costing = {
            "model": "gpt-4.1-nano",
            "prompt_tokens": prompt_tokens,
            "completion_tokens": completion_tokens,
            "total_tokens": total_tokens,
            "estimated_cost_usd": (
                (prompt_tokens * 0.00000015) +
                (completion_tokens * 0.0000006)
            )
        }

        
        
        # Remove markdown code blocks if present
        if text.startswith("```json"):
            text = text[7:]
        if text.startswith("```"):
            text = text[3:]
        if text.endswith("```"):
            text = text[:-3]
        
        text = text.strip()
        
        # Parse JSON
        plan = json.loads(text)
        
        # Validate structure
        if "summary" not in plan or "questions" not in plan:
            raise ValueError("Invalid plan structure: missing 'summary' or 'questions'")
        
        if not isinstance(plan["questions"], list):
            raise ValueError("Questions must be a list")
        
        # Accept a flexible number of questions (prefer ~10). Reject only if far from expected.
        if not (2 <= len(plan["questions"]) <= 7):
            raise ValueError(f"Expected ~10 questions (8-12 allowed), got {len(plan['questions'])}")
        
        # Attach costing directly to the plan dict for compatibility with callers
        plan["costing"] = costing
        return plan
        
    except json.JSONDecodeError as e:
        raise RuntimeError(f"Failed to parse LLM response as JSON. Response was:\n{text}\nError: {str(e)}")
    except Exception as e:
        raise RuntimeError(f"Failed to generate interview plan: {str(e)}")


def extract_resume_text(file_path: str) -> str:
    """
    Extract text content from a resume file (PDF or DOCX).
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        str: Extracted text content
        
    Raises:
        ValueError: If file format is not supported
        RuntimeError: If text extraction fails
    """
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Resume file not found: {file_path}")
    
    try:
        if file_path.lower().endswith(".pdf"):
            return _extract_pdf_text(file_path)
        elif file_path.lower().endswith(".docx"):
            return _extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file format. Only PDF and DOCX are supported. Got: {file_path}")
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from resume: {str(e)}")


def _extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if not text.strip():
            raise ValueError("PDF file appears to be empty or contains no extractable text")
        
        return text.strip()
        
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")


def _extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        text = "\n".join(text_parts)
        
        if not text.strip():
            raise ValueError("DOCX file appears to be empty or contains no extractable text")
        
        return text.strip()
        
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}")


def validate_resume_content(resume_text: str) -> bool:
    """
    Validate that the resume contains meaningful content.
    
    Args:
        resume_text: Extracted resume text
        
    Returns:
        bool: True if resume contains valid content
    """
    
    if not resume_text or len(resume_text.strip()) < 100:
        return False
    
    # Check for common resume keywords
    resume_keywords = [
        "experience", "education", "skills", "work", "job",
        "university", "college", "degree", "project", "role"
    ]
    
    text_lower = resume_text.lower()
    keyword_count = sum(1 for keyword in resume_keywords if keyword in text_lower)
    
    # Resume should contain at least 2 of these keywords
    return keyword_count >= 2
